import os
import re
import math
import tempfile
import io
import json
import hashlib
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import MultipleLocator, FuncFormatter
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL  # optional
# =========================
# JPG / Clipboard helpers
# =========================

def get_base_dir() -> Path:
    """Folder where the script is located (or cwd fallback)."""
    try:
        return Path(__file__).resolve().parent
    except Exception:
        return Path.cwd()

def auto_find_template_docx() -> str:
    """Find template in the same folder as the .py (prefer TEMPLATE.docx)."""
    base = get_base_dir()
    preferred = base / "template.docx"
    if preferred.exists():
        return str(preferred)

    # fallback: first .docx found in the same folder
    docxs = sorted(base.glob("*.docx"))
    if docxs:
        return str(docxs[0])

    return ""

def _ensure_parent_dir(path: str) -> None:
    """Ensure the folder for a file path exists."""
    Path(path).parent.mkdir(parents=True, exist_ok=True)


def save_figure_as_jpg(fig, out_path: str, dpi: int = 300, quality: int = 95):
    """
    Save matplotlib figure as JPG.
    Requires Pillow installed.
    """

    from PIL import Image  # pip install pillow

    # Save figure temporarily as PNG into memory
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, facecolor="white")
    buf.seek(0)

    # Convert PNG -> JPG
    img = Image.open(buf).convert("RGB")

    _ensure_parent_dir(out_path)
    img.save(out_path, "JPEG", quality=int(quality))

    buf.close()
def copy_figure_to_clipboard_windows(fig, dpi: int = 300):
    """
    Copy matplotlib figure to Windows clipboard.
    Requires Pillow + pywin32.
    """

    from PIL import Image
    import win32clipboard  # pip install pywin32

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, facecolor="white")
    buf.seek(0)

    img = Image.open(buf).convert("RGB")

    # Convert to BMP (clipboard format)
    output = io.BytesIO()
    img.save(output, "BMP")
    data = output.getvalue()[14:]  # remove BMP header

    output.close()
    buf.close()

    win32clipboard.OpenClipboard()
    try:
        win32clipboard.EmptyClipboard()
        win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
    finally:
        win32clipboard.CloseClipboard()
        
# =========================
# Basic / Pro unlock config
# =========================
# Default password can be overridden via environment variable OSCGEN_PRO_PASSWORD
OSCGEN_PRO_PASSWORD = os.environ.get("OSCGEN_PRO_PASSWORD", "pro123")

# =========================
# Unit helpers
# =========================
def cm_to_inches(cm: float) -> float:
    return float(cm) / 2.54
def parse_float_any(s: str, default: float) -> float:
    try:
        return float(str(s).replace(",", "."))
    except Exception:
        return default
# =========================
# Data reading / utilities
# =========================
def read_measurement_txt(txt_path: str) -> pd.DataFrame:
    """Reads TXT with header line containing 'Time [s]' and 'Value'."""
    with open(txt_path, "r", errors="ignore") as f:
        lines = f.readlines()
    start_idx = None
    for i, line in enumerate(lines):
        if "Time [s]" in line and "Value" in line:
            start_idx = i + 1
            break
    if start_idx is None:
        raise ValueError(f"Cannot find data table header 'Time [s] ... Value ...' in: {txt_path}")
    data = []
    for line in lines[start_idx:]:
        line = line.strip()
        if not line:
            continue
        parts = re.split(r"\s+", line)
        if len(parts) < 2:
            continue
        try:
            t = float(parts[0])
            v = float(parts[1])
            data.append((t, v))
        except ValueError:
            continue
    if not data:
        raise ValueError(f"No numeric data found in: {txt_path}")
    return pd.DataFrame(data, columns=["time_s", "value_v"])
def convert_to_kv(value_v: pd.Series) -> pd.Series:
    return value_v / 1000.0
def fmt_3dec_decimal_comma(x) -> str:
    try:
        if x is None or (isinstance(x, float) and pd.isna(x)):
            return ""
        return f"{float(str(x).replace(',', '.')):.3f}".replace(".", ",")
    except Exception:
        return ""
# =========================
# Missing formatting helpers (hotfix)
# =========================
def _to_float_maybe(x):
    """Parse number from Excel/string allowing decimal comma."""
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return None
    try:
        if isinstance(x, str):
            xs = x.replace(",", ".")
            m = re.search(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", xs)
            if not m:
                return None
            return float(m.group(0))
        return float(x)
    except Exception:
        return None

def fmt_voltage_to_kv(x) -> str:
    """Format a voltage-like value to kV string with decimal comma (3 dec)."""
    v = _to_float_maybe(x)
    if v is None:
        return ""
    return f"{v:.3f}".replace(".", ",")

def fmt_current_with_unit(x):
    """Return (value_str, unit) for current. Uses A / kA with decimal comma."""
    v = _to_float_maybe(x)
    if v is None:
        return ("", "")
    unit = "A"
    if abs(v) >= 1000.0:
        v = v / 1000.0
        unit = "kA"
    return (f"{v:.3f}".replace(".", ","), unit)

def fmt_time_to_us(x) -> str:
    """Format a time-like value to microseconds with decimal comma.

    Supports strings with units: ns, us (µs), ms, s.
    If no unit is present, falls back to the old heuristic:
      - If |x| < 0.01 -> treat as seconds and convert to us (x*1e6)
      - Otherwise -> treat as already in microseconds
    """
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""

    xs = str(x).strip()
    low = xs.lower().replace("µ", "u")

    v = _to_float_maybe(xs)
    if v is None:
        return ""

    # Unit-aware parsing
    if "ns" in low:
        v = v / 1000.0
    elif "us" in low or re.search(r"\bus\b", low):
        v = v
    elif "ms" in low:
        v = v * 1000.0
    elif re.search(r"\bs\b", low) and ("ms" not in low and "us" not in low and "ns" not in low):
        v = v * 1e6
    else:
        # fallback heuristic (your current behavior)
        if abs(v) < 0.01 and v != 0:
            v = v * 1e6

    return f"{v:.3f}".replace(".", ",")

def format_kv_or_mv(x) -> str:
    """Auto-format voltage: kV or MV depending on magnitude."""
    v = _to_float_maybe(x)
    if v is None:
        return ""
    if abs(v) >= 1000.0:
        return f"{v/1000.0:.3f}".replace(".", ",") + " MV"
    return f"{v:.3f}".replace(".", ",") + " kV"

def format_time_us(x) -> str:
    """Format time value as microseconds with unit 'us'."""
    s = fmt_time_to_us(x)
    return (s + " us").strip() if s else ""

def format_voltage_from_excel(x) -> str:
    """Format voltage preserving the unit implied by the Excel cell.

    Accepts either numeric or strings like '535.858 kV', '394.640 V', '1.20 MV'.
    Outputs with decimal comma.

    Rules:
    - If cell text contains 'kV' -> treat number as kV; show kV (or MV if >=1000 kV).
    - If contains 'MV' -> treat as MV; show MV.
    - If contains 'V' (but not kV/MV) -> treat as V; show V (<1000 V) else kV/MV.
    - If no unit -> fall back to kV/MV formatter (legacy behavior).
    """
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    xs = str(x)
    low = xs.lower()
    v = _to_float_maybe(x)
    if v is None:
        return ""
    # Detect units in text (order matters)
    if "mv" in low:
        return f"{v:.3f}".replace(".", ",") + " MV"
    if "kv" in low:
        # value is in kV
        if abs(v) >= 1000.0:
            return f"{(v/1000.0):.3f}".replace(".", ",") + " MV"
        return f"{v:.3f}".replace(".", ",") + " kV"
    # plain volts (avoid matching kV)
    if " v" in low or low.endswith("v"):
        return format_voltage_v_kv_mv(v)
    # unknown -> assume kV legacy
    return format_kv_or_mv(v)


def format_voltage_from_excel_preserve_unit(x, blank_if_missing_unit: bool = False) -> str:
    """Format voltage exactly in the unit written in the Excel cell.

    This is used for channel 2 (c*.txt) textbox, where the unit must match Excel.

    Accepted units in the cell text: V, kV, MV (case-insensitive). The numeric part is NOT converted.
    If no unit is present:
      - if blank_if_missing_unit=True -> return only the numeric value (no unit)
      - else -> fall back to legacy auto formatter (kV/MV)
    """
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    xs = str(x).strip()
    low = xs.lower()

    v = _to_float_maybe(x)
    if v is None:
        return ""

    # Detect unit tokens (order matters: mv, kv, then plain v)
    unit = None
    if re.search(r"\bmv\b", low):
        unit = "MV"
    elif re.search(r"\bkv\b", low):
        unit = "kV"
    elif re.search(r"(?:\bv\b|v$)", low) and not re.search(r"\bkv\b|\bmv\b", low):
        unit = "V"

    num = f"{v:.3f}".replace(".", ",")

    if unit:
        return f"{num} {unit}"
    if blank_if_missing_unit:
        return num
    # fallback (legacy)
    return format_kv_or_mv(v)


def format_voltage_v_kv_mv(x) -> str:
    """Format voltage where the *input is in volts*.

    - < 1000 V   -> V
    - < 1e6 V    -> kV
    - otherwise  -> MV

    Uses decimal comma and 3 decimals (except V: 1 decimal is usually enough, but keep 1 to stay readable).
    """
    v = _to_float_maybe(x)
    if v is None:
        return ""
    av = abs(v)
    if av < 1000.0:
        return f"{v:.1f}".replace(".", ",") + " V"
    if av < 1_000_000.0:
        return f"{(v/1000.0):.3f}".replace(".", ",") + " kV"
    return f"{(v/1_000_000.0):.3f}".replace(".", ",") + " MV"

    s = fmt_time_to_us(x)
    return (s + " us").strip() if s else ""

    try:
        if isinstance(x, str):
            x = x.replace(",", ".")
            m = re.search(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", x)
            if not m:
                return ""
            x = m.group(0)
        s = f"{float(x):.3f}"
        return s.replace(".", ",")
    except Exception:
        return ""

def fmt_tick_decimal_comma(x, _pos=None) -> str:
    """Matplotlib tick formatter using decimal comma."""
    try:
        if x is None or (isinstance(x, float) and (math.isnan(x) or math.isinf(x))):
            return ""
        # Prefer integer formatting when very close to an integer
        if abs(x - round(x)) < 1e-9:
            s = str(int(round(x)))
        else:
            s = f"{x:.3f}".rstrip("0").rstrip(".")
        return s.replace(".", ",")
    except Exception:
        try:
            return str(x).replace(".", ",")
        except Exception:
            return ""

def normalize_shape(s: str) -> str:
    """Normalize Shape strings (trim and collapse multiple spaces)."""
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s
def is_current_type(type_str: str) -> bool:
    return normalize_shape(type_str).lower() == "current"


# =========================
# Excel metadata loading
# =========================
def load_excel_metadata(excel_path: str) -> pd.DataFrame:
    """Load required metadata columns from Excel.

    NOTE (hardcoded positional mode):
    This version supports loading an Excel sheet even when it has NO header row.
    In that case, columns are taken by fixed position (1-based Excel columns):

      1  -> No
      2  -> Shape
      3  -> Type
      4  -> PkMax
      5  -> Min
      6  -> T1
      7  -> Tp
      8  -> T2
      9  -> Td
      10 -> Tc
      11 -> T0

    The normal mode (with headers) remains supported for convenience.
    """

    # -------------------------
    # CSV SUPPORT (minimal add)
    # -------------------------
    ext = Path(excel_path).suffix.lower()

    if ext == ".csv":
        # 1) Read CSV with delimiter sniffing
        df = pd.read_csv(excel_path, sep=None, engine="python", encoding="utf-8-sig")

        # 2) If it ended up as ONE column, split it manually
        if df.shape[1] == 1:
            col0 = df.columns[0]
            s = df[col0].astype(str)

            split = s.str.split(",", expand=True)

            if split.shape[1] == 1:
                split = s.str.split(";", expand=True)

            split.columns = split.iloc[0].astype(str).tolist()
            df = split.iloc[1:].reset_index(drop=True)

        cols = {str(c).strip(): c for c in df.columns}

        def find_col(candidates: list[str]):
            for cand in candidates:
                for k in cols.keys():
                    if cand.lower() in k.lower():
                        return cols[k]
            return None

        col_no = find_col(["No"])
        col_shape = find_col(["Shape"])
        col_type = find_col(["Type"])
        col_pkmax = find_col(["Pk;Max", "PkMax", "Max"])
        col_min = find_col(["Min"])
        col_t1 = find_col(["T1"])
        col_tp = find_col(["Tp"])
        col_t2 = find_col(["T2"])
        col_td = find_col(["Td"])
        col_tc = find_col(["Tc"])
        col_t0 = find_col(["To", "T0"])

        missing = [n for n, c in [
            ("No", col_no), ("Shape", col_shape), ("Type", col_type),
            ("PkMax", col_pkmax), ("Min", col_min),
            ("T1", col_t1), ("Tp", col_tp), ("T2", col_t2),
            ("Td", col_td), ("Tc", col_tc), ("T0", col_t0)
        ] if c is None]

        if missing:
            raise ValueError("CSV loaded but missing required columns: " + ", ".join(missing))

        meta = df[[col_no, col_shape, col_type,
                   col_pkmax, col_min,
                   col_t1, col_tp, col_t2,
                   col_td, col_tc, col_t0]].copy()

        meta.columns = ["No", "Shape", "Type", "PkMax", "Min", "T1", "Tp", "T2", "Td", "Tc", "T0"]
        meta = meta.dropna(subset=["No"]).reset_index(drop=True)
        return meta

    # -------------------------
    # Attempt A: normal workbook with header row (existing behavior)
    # -------------------------
    try:
        df = pd.read_excel(excel_path, engine="openpyxl")
        cols = {str(c).strip(): c for c in df.columns}

        def find_col(candidates: list[str]):
            for cand in candidates:
                for k in cols.keys():
                    if cand.lower() in k.lower():
                        return cols[k]
            return None

        col_no = find_col(["No"])  # "No." or "No"
        col_shape = find_col(["Shape"])
        col_type = find_col(["Type"])

        # Amplitudes
        col_pkmax = find_col(["Pk;Max", "Pk; Max", "PkMax", "Max"])
        col_min = find_col(["Min"])

        # Times
        col_t1 = find_col(["T1"])
        col_tp = find_col(["Tp"])
        col_t2 = find_col(["T2"])
        col_td = find_col(["Td"])
        col_tc = find_col(["Tc"])
        col_t0 = find_col(["To", "T0"])

        missing = [
            name
            for name, col in [
                ("No", col_no),
                ("Shape", col_shape),
                ("Type", col_type),
                ("PkMax", col_pkmax),
                ("Min", col_min),
                ("T1", col_t1),
                ("Tp", col_tp),
                ("T2", col_t2),
                ("Td", col_td),
                ("Tc", col_tc),
                ("T0", col_t0),
            ]
            if col is None
        ]

        if not missing:
            meta = df[
                [col_no, col_shape, col_type, col_pkmax, col_min, col_t1, col_tp, col_t2, col_td, col_tc, col_t0]
            ].copy()
            meta.columns = ["No", "Shape", "Type", "PkMax", "Min", "T1", "Tp", "T2", "Td", "Tc", "T0"]
            meta = meta.dropna(subset=["No"]).reset_index(drop=True)
            return meta

        # If headers exist but required columns are not found, fall back to positional.
    except Exception:
        # If normal read fails, fall back to positional.
        pass

    # -------------------------
    # Attempt B: NO header row -> fixed positions
    # -------------------------
    df2 = pd.read_excel(excel_path, engine="openpyxl", header=None)

    if df2.shape[1] < 11:
        raise ValueError(
            f"Excel without header detected, but it has only {df2.shape[1]} columns. "
            "Expected at least 11 columns in fixed order: No, Shape, Type, PkMax, Min, T1, Tp, T2, Td, Tc, T0."
        )

    meta = df2.iloc[:, 0:11].copy()
    meta.columns = ["No", "Shape", "Type", "PkMax", "Min", "T1", "Tp", "T2", "Td", "Tc", "T0"]
    meta = meta.dropna(subset=["No"]).reset_index(drop=True)
    return meta

def fmt_kv_from_v(x) -> str:
    """Back-compat wrapper: format voltage-like values to kV."""
    return fmt_voltage_to_kv(x)


def fmt_a(x) -> str:
    """Back-compat wrapper: format current-like values (numeric part only)."""
    v, _u = fmt_current_with_unit(x)
    return v


    try:
        return fmt_3dec_decimal_comma(x)
    except Exception:
        return ""
def fmt_us_from_s(x) -> str:
    """Back-compat wrapper: format time-like values to microseconds."""
    return fmt_time_to_us(x)

def format_current_auto(value_a: float | None) -> str:
    """Format current value with automatic A/kA unit and decimal comma."""
    if value_a is None or (isinstance(value_a, float) and pd.isna(value_a)):
        return ""
    try:
        v = float(value_a)
    except Exception:
        return ""
    unit = "A"
    if abs(v) >= 1000.0:
        v = v / 1000.0
        unit = "kA"
    s = f"{v:.3f}".replace(".", ",")
    return f"{s} {unit}"


def build_textbox_lines(meta: dict, file_prefix: str | None = None, textbox_mode_c: str = "Auto (from Excel)") -> list[str]:
    """Build textbox content.

    - For v-files: always use Excel-driven rules (Shape + Type).
    - For c-files: if textbox_mode_c != 'Auto (from Excel)', use the selected template.
    """
    no = meta.get("No", "")
    shape_raw = str(meta.get("Shape", ""))
    shape_norm = normalize_shape(shape_raw)
    typ = str(meta.get("Type", ""))
    cur = is_current_type(typ)

    pkmax = meta.get("PkMax")
    mn = meta.get("Min")
    t1 = meta.get("T1")
    tp = meta.get("Tp")
    t2 = meta.get("T2")
    td = meta.get("Td")
    tc = meta.get("Tc")
    t0 = meta.get("T0")


    # Voltage formatter:
    # - v-files: keep existing Excel-driven formatting (with units if present)
    # - c-files: show the SAME unit as written in Excel; if Excel has no unit -> show only the number (no unit)
    _fp = (file_prefix or "").lower().strip()
    if _fp == "c":
        fmt_volt = lambda x: format_voltage_from_excel_preserve_unit(x, blank_if_missing_unit=True)
    else:
        fmt_volt = format_voltage_from_excel

    def add_if(out: list[str], label: str, value_str: str):
        value_str = (value_str or "").strip()
        if value_str:
            out.append(f"{label}:  {value_str}")

    # Apply user-selected textbox template for c-files (channel 2)
    mode = (textbox_mode_c or "Auto (from Excel)").strip()
    if (file_prefix or "").lower() == "c" and mode != "Auto (from Excel)":
        # Minimal MATLAB-style template:
        # TEXT_BOX={["No.",num2str(headertext{1})];["Upk:  ",headertext{6}];};
        if mode == "Default":
            return [
                f"No.{no}",
                f"Upk:  {(fmt_volt(pkmax) if (file_prefix or '').lower() == 'c' else fmt_volt(pkmax))}",
            ]
        if mode.startswith("LI full") and "Ipk max/min" in mode:
            return [
                f"No.{no}",
                "LI full",
                f"Ipk max:  {format_current_auto(pkmax)}",
                f"Ipk min:  {format_current_auto(mn)}",
            ]
        if mode.startswith("LI full") and "(Ipk)" in mode:
            return [
                f"No.{no}",
                "LI full",
                f"Ipk:  {format_current_auto(pkmax)}",
            ]
        if mode.startswith("LI full") and "Upk" in mode:
            return [
                f"No.{no}",
                "LI full",
                f"Upk:  {fmt_volt(pkmax)}",
                f"T1:  {format_time_us(t1)}",
                f"T2:  {format_time_us(t2)}",
            ]

        if mode.startswith("LI tailchopped") and "(Ipk)" in mode:
            return [
                f"No.{no}",
                "LI tailchopped",
                f"Ipk:  {format_current_auto(pkmax)}",
            ]
        if mode.startswith("LI tailchopped") and "Upk" in mode:
            # MATLAB style: T2 is shown with '--'
            t2s = format_time_us(t2)
            t2line = f"T2:-- {t2s}".rstrip() if t2s else "T2:--"
            return [
                f"No.{no}",
                "LI tailchopped",
                f"Upk min:  {fmt_volt(mn)}",
                f"Upk max:  {fmt_volt(pkmax)}",
                f"T1:  {format_time_us(t1)}",
                t2line,
                f"Tc:  {format_time_us(tc)}",
            ]

        if mode.startswith("SI IEC 60060") and "(Ipk" in mode:
            return [
                f"No.{no}",
                "SI IEC 60060",
                f"Ipk:  {format_current_auto(pkmax)}",
                f"Tp:  {format_time_us(tp)}",
                f"T2:  {format_time_us(t2)}",
            ]
        if mode.startswith("SI IEC 60060") and "Upk" in mode:
            return [
                f"No.{no}",
                "SI IEC 60060",
                f"Upk:  {fmt_volt(pkmax)}",
                f"Tp:  {format_time_us(tp)}",
                f"T2:  {format_time_us(t2)}",
            ]

        if mode.startswith("SI IEC 60076") and "Ipk max/min" in mode:
            return [
                f"No.{no}",
                "SI IEC 60076",
                f"Ipk max:  {format_current_auto(pkmax)}",
                f"Ipk min:  {format_current_auto(mn)}",
            ]
        if mode.startswith("SI IEC 60076") and "Upk" in mode:
            return [
                f"No.{no}",
                "SI IEC 60076",
                f"Upk:  {fmt_volt(pkmax)}",
                f"T1:  {format_time_us(t1)}",
                f"Td:  {format_time_us(td)}",
                f"T0:  {format_time_us(t0)}",
            ]

        # Fallback: show a minimal box
        return [f"No.{no}", f"Upk:  {fmt_volt(pkmax)}"]

    # Default behavior (Excel-driven rules, for v-files and auto mode for c)
    # Match MATLAB behavior using normalized shape string:
    # Default behavior (Excel-driven rules, for v-files and auto mode for c)
    # Match MATLAB behavior using normalized shape string:
    if shape_norm == "LI full":
        out = [f"No.{no}", "LI full"]

        if cur:
            add_if(out, "Ipk", format_current_auto(pkmax))
            return out

        add_if(out, "Upk", fmt_volt(pkmax))
        add_if(out, "T1", format_time_us(t1))
        add_if(out, "T2", format_time_us(t2))
        return out
    if shape_norm == "LI frontchopped":
        out = [f"No.{no}", "LI frontchopped"]

        if cur:
            add_if(out, "Ipk", format_current_auto(pkmax))
            add_if(out, "T1", format_time_us(t1))
            add_if(out, "Tc", format_time_us(tc))
            return out

        # Voltage variant
        add_if(out, "Upk min", fmt_volt(mn))
        add_if(out, "Upk max", fmt_volt(pkmax))
        add_if(out, "T1", format_time_us(t1))
        add_if(out, "Tc", format_time_us(tc))
        return out
    if shape_norm == "LI tailchopped":
        if cur:
            return [f"No.{no}", "LI tailchopped", f"Ipk:  {format_current_auto(pkmax)}"]
        t2s = format_time_us(t2)
        t2line = f"T2:-- {t2s}".rstrip() if t2s else "T2:--"
        out = [f"No.{no}", "LI tailchopped"]
        add_if(out, "Upk min", fmt_volt(mn))
        add_if(out, "Upk max", fmt_volt(pkmax))
        add_if(out, "T1", format_time_us(t1))

        t2s = (format_time_us(t2) or "").strip()
        if t2s:
            out.append(f"T2:-- {t2s}")

        add_if(out, "Tc", format_time_us(tc))
        return out
    if shape_norm == "SI IEC 60060":
        if cur:
            return [f"No.{no}", "SI IEC 60060", f"Ipk:  {format_current_auto(pkmax)}", f"Tp:  {format_time_us(tp)}", f"T2:  {format_time_us(t2)}"]
        out = [f"No.{no}", "SI IEC 60060"]
        add_if(out, "Upk", fmt_volt(pkmax))
        add_if(out, "Tp", format_time_us(tp))
        add_if(out, "T2", format_time_us(t2))
        return out
    if shape_norm == "SI IEC 60076":
        if cur:
            # Best effort: if it's current and this shape is used, show max/min like MATLAB's spaced variant.
            return [f"No.{no}", "SI IEC 60076", f"Ipk max:  {format_current_auto(pkmax)}", f"Ipk min:  {format_current_auto(mn)}"]
        out = [f"No.{no}", "SI IEC 60076"]
        add_if(out, "Upk", fmt_volt(pkmax))
        add_if(out, "T1", format_time_us(t1))
        add_if(out, "Td", format_time_us(td))
        add_if(out, "T0", format_time_us(t0))
        return out

    # Fallback
    if cur:
        return [f"No.{no}", f"Ipk:  {format_current_auto(pkmax)}"]
    return [f"No.{no}", f"Upk:  {fmt_volt(pkmax)}"]




def choose_nice_step(data_min: float, data_max: float, target_ticks: int = 8) -> float:
    """Pick a 'nice' major tick step for Y axis.

    Uses 1/2/5 * 10^n stepping similar to MATLAB 'nice' ticks.
    """
    span = abs(float(data_max) - float(data_min))
    if span == 0:
        return 1.0
    raw = span / float(target_ticks)
    exp = 10 ** (math.floor(math.log10(raw)))
    frac = raw / exp
    if frac <= 1:
        nice = 1
    elif frac <= 2:
        nice = 2
    elif frac <= 5:
        nice = 5
    else:
        nice = 10
    return float(nice * exp)

def grid_params(grid_scale: str):
    """
    Returns:
    x_minor_step_us, y_minor_div, major_lw, minor_lw
    """
    gs = (grid_scale or "Normal").lower()
    if gs == "fine":
        # keep denser minor ticks but don't draw minor grid by default
        return 2.0, 4, 0.85, 0.0
    if gs == "coarse":
        return 5.0, 2, 0.85, 0.0  # no minor grid
    # "Normal": keep minor ticks (for easier reading) but DO NOT draw minor grid.
    # This avoids the extra grid lines on unlabeled minor ticks.
    return 5.0, 2, 0.85, 0.0
def draw_oscillogram_on_axes(
    ax,
    df: pd.DataFrame,
    serial_number: str,
    divider_str: str,
    meta: dict,
    xmax_us: float,
    y_mode: str,
    y_step_kv: float | None,
    grid_scale: str,
    file_prefix: str | None = None,
    divider1_str: str = "",
    divider2_str: str = "",
    show_matlab_header: bool = True,
    textbox_mode_c: str = "Auto (from Excel)",
    grid_major_lw: float | None = None,
    grid_minor_lw: float | None = None,
    plot_frame_lw: float = 1.2,
    textbox_frame_lw: float = 1.0,
    tick_major_lw: float = 1.0,
    tick_minor_lw: float = 0.8,
    font_label: int = 6,
    font_tick: int = 6,
    font_header: int = 6,
    font_textbox: int = 6,
    y_scale: float = 1.0,
    textbox_pos_mode: str = "auto",
):
    t_us = df["time_s"] * 1e6

    # df["value_v"] is in VOLTS. Apply optional scaling, then choose display unit automatically: V / kV / MV.
    try:
        _ys = float(y_scale)
    except Exception:
        _ys = 1.0
    y_v = df["value_v"] * _ys
    peak_v = float(max(abs(float(y_v.min())), abs(float(y_v.max()))))

    if peak_v < 1000.0:
        y_plot = y_v
        y_unit_label = "V"
    elif peak_v < 1_000_000.0:
        y_plot = y_v / 1000.0
        y_unit_label = "kV"
    else:
        y_plot = y_v / 1_000_000.0
        y_unit_label = "MV"

    use_mv = (y_unit_label == "MV")

    ax.clear()
    # Curve (thin)
    ax.plot(t_us, y_plot, linewidth=1.0, color="green")
    # Font sizes (a bit smaller to better match typical Word report look)
    ax.set_xlabel("TIME (us)", fontsize=font_label)
    ax.set_ylabel(f"VOLTAGE ({y_unit_label})", fontsize=font_label)
    ax.tick_params(labelsize=font_tick)
    ax.tick_params(which="major", width=float(tick_major_lw))
    ax.tick_params(which="minor", width=float(tick_minor_lw))
    # X
    # Special rule: tailchopped impulses should use a fixed window:
    #   Xmax = 10 us, with 10% of that in front of zero.
    # i.e. [-1 us .. 10 us]
    shape_norm = normalize_shape(str(meta.get("Shape", ""))).lower()

    if "tailchopped" in shape_norm:
        xmin_us = -1.0
        xmax_plot = 10.0
    else:
        xmax_plot = float(xmax_us)
        # For 'c' files, use 10% pretrigger (as requested)
        if (file_prefix or "").lower() == "c":
            xmin_us = -0.10 * xmax_plot
        else:
            xmin_us = -5.0

    ax.set_xlim(xmin_us, xmax_plot)

    # Grid params (shared)
    x_minor_step_us, y_minor_div, major_lw, minor_lw = grid_params(grid_scale)
    if grid_major_lw is not None:
        major_lw = float(grid_major_lw)
    if grid_minor_lw is not None:
        minor_lw = float(grid_minor_lw)

    # X ticks:
    # - tailchopped: fixed [-1..10] us with 1 us major ticks
    # - otherwise (v & c): 10.5 major divisions total
    #   -> 0.5 division left of zero, 10 divisions right of zero
    if "tailchopped" in shape_norm:
        ax.xaxis.set_major_locator(MultipleLocator(1))
        ax.xaxis.set_minor_locator(MultipleLocator(0.2))
    else:
        # Major step so that right side has exactly 10 divisions up to xmax_plot
        major_step_x = float(xmax_plot) / 10.0

        # Shift so that zero is not on the border:
        # 0.5 major division to the left, 10 to the right
        xmin_us = -0.5 * major_step_x
        xmax_plot = 10.0 * major_step_x
        ax.set_xlim(xmin_us, xmax_plot)

        ax.xaxis.set_major_locator(MultipleLocator(major_step_x))
        ax.xaxis.set_minor_locator(MultipleLocator(major_step_x / 5.0))

    # Decimal comma on ticks (both axes)
    ax.xaxis.set_major_formatter(FuncFormatter(fmt_tick_decimal_comma))
    ax.yaxis.set_major_formatter(FuncFormatter(fmt_tick_decimal_comma))
    # Y
    y_min = float(y_plot.min())
    y_max = float(y_plot.max())
    if y_mode == "manual":
        major_step = float(y_step_kv) if (y_step_kv is not None and y_step_kv > 0) else 20.0
        if use_mv:
            major_step = major_step / 1000.0
    else:
        major_step = choose_nice_step(y_min, y_max, target_ticks=8)
    y0 = major_step * math.floor(y_min / major_step)
    y1 = major_step * math.ceil(y_max / major_step)
    if y1 == y0:
        y1 += major_step
    ax.set_ylim(y0, y1)
    ax.yaxis.set_major_locator(MultipleLocator(major_step))
    if y_minor_div and y_minor_div > 1:
        ax.yaxis.set_minor_locator(MultipleLocator(major_step / y_minor_div))
    # Grid (make it crisp like your example)
    ax.grid(True, which="major", linestyle="--", linewidth=major_lw, color="black")
    if minor_lw > 0:
        ax.grid(True, which="minor", linestyle="--", linewidth=minor_lw, color="black")
        # --- Solid axes at x=0 and y=0 (over the dashed grid) ---
    x0, x1 = ax.get_xlim()
    y0, y1 = ax.get_ylim()
    # draw only if 0 is inside the visible range
    if x0 <= 0 <= x1:
        ax.axvline(0, color="black", linewidth=major_lw, linestyle="-", zorder=3)
    if y0 <= 0 <= y1:
        ax.axhline(0, color="black", linewidth=major_lw, linestyle="-", zorder=3)
    # Frame & ticks a bit stronger
    for spine in ax.spines.values():
        spine.set_linewidth(float(plot_frame_lw))
        spine.set_color("black")
    ax.tick_params(axis="both", which="major", width=float(tick_major_lw), length=3, colors="black")
    ax.tick_params(axis="both", which="minor", width=float(tick_minor_lw), length=2, colors="black")
    # Headers
    # MATLAB-style header: TWO anchored texts (as requested):
    #  - LEFT: ' Serial number: ...' (starts at left edge with a single leading space)
    #  - RIGHT: 'CHANNEL X (VOLTAGE)   Divider: ...' anchored to right edge
    if show_matlab_header and file_prefix in ("v", "c"):
        if file_prefix == "v":
            right_text = f"CHANNEL 1 (VOLTAGE)   Divider: {divider1_str}"
        else:
            right_text = f"CHANNEL 2 (VOLTAGE)   Divider: {divider2_str}"

        ax.text(
            -0.02, 1.04,
            f"Serial number: {serial_number}",
            transform=ax.transAxes,
            ha="left", va="bottom",
            fontsize=font_header,
            color="black",
            clip_on=False,
        )
        ax.text(
            1.0, 1.04,
            right_text,
            transform=ax.transAxes,
            ha="right", va="bottom",
            fontsize=font_header,
            color="black",
            clip_on=False,
        )
    else:
        # Fallback (should rarely happen): keep a simple serial line.
        ax.text(
            0.0, 1.08,
            f"Serial number: {serial_number}",
            transform=ax.transAxes,
            ha="left", va="bottom",
            fontsize=font_header,
            color="black",
            clip_on=False,
        )

        ax.text(
            1.0, 1.02,
            f"Divider: {divider_str} V/V",
            transform=ax.transAxes,
            ha="right", va="bottom",
            fontsize=font_header,
            color="black"
        )
    # Info box
    box_lines = build_textbox_lines(meta, file_prefix=file_prefix, textbox_mode_c=textbox_mode_c)
    box_text = "\n".join(box_lines)
    # Info box (position can be auto based on waveform polarity)
    mode = (textbox_pos_mode or "auto").strip().lower()
    if mode in ("auto", "a"):
        # Heuristic (INVERTED by request):
        # - If waveform is mostly positive -> put textbox near TOP-RIGHT
        # - If waveform is mostly negative -> put textbox near BOTTOM-RIGHT
        # - If it crosses zero -> use the dominant absolute peak
        if y_max <= 0:
            corner = "bottom-right"
        elif y_min >= 0:
            corner = "top-right"
        else:
            corner = "top-right" if abs(y_max) >= abs(y_min) else "bottom-right"
    elif mode in ("tr", "top", "top-right", "topright"):
        corner = "top-right"
    else:
        corner = "bottom-right"
    if corner == "top-right":
        tx, ty, ha, va = 0.98, 0.965, "right", "top"
    else:
        tx, ty, ha, va = 0.98, 0.04, "right", "bottom"
    ax.text(
        tx, ty, box_text,
        transform=ax.transAxes,
        ha=ha, va=va,
        multialignment="left",
        bbox=dict(
            boxstyle="square,pad=0.35",
            facecolor="white",
            edgecolor="black",
            linewidth=float(textbox_frame_lw),
        ),
        fontsize=font_textbox,
        color="black",
        clip_on=True,
    )
def make_oscillogram_png(
    txt_path: str,
    out_png: str,
    serial_number: str,
    divider_str: str,
    meta: dict,
    xmax_us: float,
    y_mode: str,
    y_step_kv: float | None,
    grid_scale: str,
    out_width_cm: float,
    out_height_cm: float,
    out_dpi: int,
    grid_major_lw: float | None = None,
    grid_minor_lw: float | None = None,
    plot_frame_lw: float = 1.2,
    textbox_frame_lw: float = 1.0,
    tick_major_lw: float = 1.0,
    tick_minor_lw: float = 0.8,
    font_label: int = 6,
    font_tick: int = 6,
    font_header: int = 6,
    font_textbox: int = 6,
    textbox_mode_c: str = "Auto (from Excel)",
    textbox_pos_mode: str = "auto",
    y_scale: float = 1.0,
    file_prefix: str | None = None,
    divider1_str: str = "",
    divider2_str: str = "",
    show_matlab_header: bool = True,
    margin_left: float = 0.12,
    margin_right: float = 0.93,
    margin_bottom: float = 0.22,
    margin_top: float = 0.88,
    use_tight_bbox: bool = False,
    pad_inches: float = 0.02,
):
    """
    Key for sharpness:
    - figure size matches final Word size (cm)
    - DPI high (300–600)
    """
    df = read_measurement_txt(txt_path)
    fig_w_in = cm_to_inches(out_width_cm)
    fig_h_in = cm_to_inches(out_height_cm)
    fig = plt.figure(figsize=(fig_w_in, fig_h_in), dpi=out_dpi)
    ax = fig.add_subplot(111)
    draw_oscillogram_on_axes(
        ax, df, serial_number, divider_str, meta,
        file_prefix=file_prefix,
        divider1_str=divider1_str,
        divider2_str=divider2_str,
        show_matlab_header=show_matlab_header,
        textbox_mode_c=textbox_mode_c,
        xmax_us=xmax_us, y_mode=y_mode, y_step_kv=y_step_kv, grid_scale=grid_scale,
        grid_major_lw=grid_major_lw,
        grid_minor_lw=grid_minor_lw,
        plot_frame_lw=plot_frame_lw,
        textbox_frame_lw=textbox_frame_lw,
        tick_major_lw=tick_major_lw,
        tick_minor_lw=tick_minor_lw,
        font_label=font_label,
        font_tick=font_tick,
        font_header=font_header,
        font_textbox=font_textbox,
        textbox_pos_mode=textbox_pos_mode,
        y_scale=y_scale,
    )
    # Margins around the plot area (MATLAB-like defaults)
    # Values are relative in [0..1] with respect to the figure.
    fig.subplots_adjust(
        left=float(margin_left),
        right=float(margin_right),
        bottom=float(margin_bottom),
        top=float(margin_top),
    )
    # Save exact size. Optional tight bbox if user wants automatic cropping.
    save_kwargs = dict(dpi=out_dpi, facecolor="white")
    if use_tight_bbox:
        save_kwargs.update({"bbox_inches": "tight", "pad_inches": float(pad_inches)})
    fig.savefig(out_png, **save_kwargs)
    plt.close(fig)
# =========================
# Word insertion (STRICT TEMPLATE)
# =========================
def clear_cell_keep_format(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if len(cell.paragraphs) > 1:
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    if not cell.paragraphs:
        cell.add_paragraph("")

def _parse_textbox_lines_to_row(lines: list[str]) -> dict:
    """Parse textbox lines (as rendered on the oscillogram) into a flat dict.

    Rules:
    - Lines like 'Upk:  535,858 kV' become {'Upk': '535,858 kV'}
    - 'No.5' becomes {'No.': '5'}
    - The first non-empty line without ':' after 'No.' is treated as 'Shape'
    - Any other non-empty no-colon lines are stored as 'Info' (concatenated)
    """
    row: dict[str, str] = {}
    info_extra: list[str] = []
    for raw in (lines or []):
        s = (raw or "").strip()
        if not s:
            continue
        if ":" in s:
            k, v = s.split(":", 1)
            k = k.strip()
            v = v.strip()
            if k.lower().startswith("upk"):
                v = v.replace("kV", "").replace("MV", "").strip()
            if k:
                row[k] = v
            continue
        low = s.lower()
        if low.startswith("no."):
            row["No."] = s[3:].strip()
            continue
        if "Shape" not in row:
            row["Shape"] = s
        else:
            info_extra.append(s)
    if info_extra:
        row["Info"] = " / ".join(info_extra)
    return row

def _preferred_table_columns(all_keys: list[str]) -> list[str]:
    preferred = [
        "No.",
        "Shape",
        "Upk",
        "Upk min",
        "Upk max",
        "Ipk",
        "Ipk max",
        "Ipk min",
        "T1",
        "Tp",
        "T2",
        "Td",
        "Tc",
        "T0",
        "Info",
    ]
    out = []
    for k in preferred:
        if k in all_keys and k not in out:
            out.append(k)
    # add any remaining keys (stable order)
    for k in all_keys:
        if k not in out:
            out.append(k)
    return out

def _append_textbox_values_table(doc: Document, table_rows: list[dict]) -> None:
    """Append a Word table with textbox values for v* oscillograms."""
    if not table_rows:
        return
    # Collect all keys
    keys = []
    for r in table_rows:
        for k in r.keys():
            if k not in keys:
                keys.append(k)
    cols = _preferred_table_columns(keys)
    if not cols:
        return

    doc.add_paragraph("")
    doc.add_paragraph("Textbox values – voltage channel (v*)")
    t = doc.add_table(rows=1, cols=len(cols))
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    hdr = t.rows[0].cells
    for j, k in enumerate(cols):
        hdr[j].text = k

    for r in table_rows:
        cells = t.add_row().cells
        for j, k in enumerate(cols):
            cells[j].text = str(r.get(k, ""))

def insert_images_into_template_strict(
    template_docx: str,
    images: list[str],
    output_docx: str,
    width_cm: float,
    height_cm: float,
    lock_aspect: bool,
    peaks_table_rows: list[dict] | None = None,
):
    doc = Document(template_docx)
    if not doc.tables:
        raise ValueError("The Word template contains no tables.")
    tbl = doc.tables[0]
    rows = len(tbl.rows)
    cols = len(tbl.columns)
    capacity = rows * cols
    if len(images) > capacity:
        raise ValueError(f"Template capacity is {capacity} images, but you generated {len(images)}.")
    w_in = cm_to_inches(width_cm)
    h_in = cm_to_inches(height_cm)
    idx = 0
    for r in range(rows):
        for c in range(cols):
            if idx >= len(images):
                break
            cell = tbl.cell(r, c)
            clear_cell_keep_format(cell)

            # optional vertical centering
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            tcPr = cell._tc.get_or_add_tcPr()

            # remove existing tcMar if present
            for el in tcPr.findall(qn("w:tcMar")):
                tcPr.remove(el)

            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "0")
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
            p.paragraph_format.line_spacing = 1

            run = p.add_run()
            if lock_aspect:
                run.add_picture(images[idx], width=Inches(w_in))
            else:
                run.add_picture(images[idx], width=Inches(w_in), height=Inches(h_in))
            idx += 1
        if idx >= len(images):
            break
    # Optional: append a summary table with textbox values (for v-oscillograms)
    if peaks_table_rows:
        try:
            _append_textbox_values_table(doc, peaks_table_rows)
        except Exception:
            # Don't fail the whole export if table fails
            pass
    doc.save(output_docx)
def export_strict_template_multipart(
    template_docx: str,
    images: list[str],
    output_base_docx: str,
    width_cm: float,
    height_cm: float,
    lock_aspect: bool,
):
    probe = Document(template_docx)
    if not probe.tables:
        raise ValueError("The Word template contains no tables.")
    tbl = probe.tables[0]
    capacity = len(tbl.rows) * len(tbl.columns)
    base = Path(output_base_docx)
    stem = base.stem
    suffix = base.suffix
    parts = []
    for start in range(0, len(images), capacity):
        part_images = images[start:start + capacity]
        part_idx = (start // capacity) + 1
        out_part = str(base.with_name(f"{stem}_part{part_idx:02d}{suffix}"))
        insert_images_into_template_strict(
            template_docx, part_images, out_part,
            width_cm=width_cm, height_cm=height_cm, lock_aspect=lock_aspect
        )
        parts.append(out_part)
    return parts, capacity
# =========================
# Preview Panel (Tab 2)
# =========================
class OscPreviewPanel(ttk.Frame):
    def __init__(self, parent, app_get_state_callable):
        super().__init__(parent)
        self.get_state = app_get_state_callable
        self.current_index = 0
        left = ttk.Frame(self)
        left.pack(side="left", fill="y", padx=8, pady=8)
        ttk.Label(left, text="Oscillograms").pack(anchor="w")
        self.listbox = tk.Listbox(left, height=22, width=30)
        self.listbox.pack(fill="y", expand=False)
        right = ttk.Frame(self)
        right.pack(side="left", fill="both", expand=True, padx=8, pady=8)
        self.fig = plt.Figure(figsize=(9, 5), dpi=100)
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvasTkAgg(self.fig, master=right)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)
        ctrl = ttk.Frame(right)
        ctrl.pack(fill="x", pady=8)
        ttk.Label(ctrl, text="X max (us):").grid(row=0, column=0, sticky="w")
        self.xmax_var = tk.StringVar(value="100")
        self.xmax_entry = ttk.Entry(ctrl, textvariable=self.xmax_var, width=10)
        self.xmax_entry.grid(row=0, column=1, sticky="w", padx=(6, 18))
        self.y_mode_var = tk.StringVar(value="auto")
        ttk.Radiobutton(ctrl, text="Auto Y", variable=self.y_mode_var, value="auto",
                        command=self.on_controls_changed).grid(row=0, column=2, sticky="w", padx=(0, 12))
        ttk.Radiobutton(ctrl, text="Manual step (kV)", variable=self.y_mode_var, value="manual",
                        command=self.on_controls_changed).grid(row=0, column=3, sticky="w")
        self.y_step_var = tk.StringVar(value="50")
        self.y_step_entry = ttk.Entry(ctrl, textvariable=self.y_step_var, width=8)
        self.y_step_entry.grid(row=0, column=4, sticky="w", padx=(6, 18))
        ttk.Label(ctrl, text="Scale:").grid(row=0, column=5, sticky="w")
        self.scale_var = tk.StringVar(value="1.0")
        self.scale_entry = ttk.Entry(ctrl, textvariable=self.scale_var, width=8)
        self.scale_entry.grid(row=0, column=6, sticky="w", padx=(6, 0))
        for w in (self.xmax_entry, self.y_step_entry, self.scale_entry):
            w.bind("<Return>", lambda e: self.on_controls_changed())
            w.bind("<FocusOut>", lambda e: self.on_controls_changed())
        btns = ttk.Frame(right)
        btns.pack(fill="x", pady=(4, 0))
        ttk.Button(btns, text="Apply current to ALL", command=self.apply_current_to_all).pack(side="left")
        ttk.Button(btns, text="Reset selected", command=self.reset_selected).pack(side="left", padx=8)
        # Right-side actions (Step 2): save/copy the currently shown oscillogram
        ttk.Button(btns, text="Save JPG…", command=self.save_current_as_jpg).pack(side="right")
        ttk.Button(btns, text="Copy oscillogram", command=self.copy_current_to_clipboard).pack(side="right", padx=8)
        ttk.Button(btns, text="Refresh", command=self.refresh_current).pack(side="right", padx=8)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)
    def load_list(self):
        st = self.get_state()
        self.listbox.delete(0, "end")
        if not st or not st.get("items"):
            return
        for i, it in enumerate(st["items"]):
            meta = it.get("meta", {}) or {}
            no = meta.get("No", it.get("No", ""))
            shape = normalize_shape(str(meta.get("Shape", "")))
            pref = it.get("prefix", "?")
            self.listbox.insert("end", f"{pref}{no}  No.{no}  {shape}")
        self.current_index = 0
        self.listbox.selection_set(0)
        self.listbox.activate(0)
        self.load_controls_from_settings(0)
        self.render_preview(0)
    def load_controls_from_settings(self, idx: int):
        st = self.get_state()
        settings = st["settings"][idx]
        self.xmax_var.set(str(settings["xmax_us"]))
        self.y_mode_var.set(settings["y_mode"])
        self.y_step_var.set("" if settings["y_step_kv"] is None else str(settings["y_step_kv"]))
        self.scale_var.set(str(settings.get("scale_factor", 1.0)))
    def save_controls_to_settings(self, idx: int):
        st = self.get_state()
        settings = st["settings"][idx]
        settings["xmax_us"] = parse_float_any(self.xmax_var.get(), settings["xmax_us"])
        settings["y_mode"] = self.y_mode_var.get()
        if settings["y_mode"] == "manual":
            settings["y_step_kv"] = parse_float_any(self.y_step_var.get(), 50.0)
        else:
            settings["y_step_kv"] = None
        settings["scale_factor"] = parse_float_any(self.scale_var.get(), settings.get("scale_factor", 1.0))
    def on_select(self, _evt=None):
        sel = self.listbox.curselection()
        if not sel:
            return
        prev_idx = getattr(self, "current_index", 0)
        try:
            self.save_controls_to_settings(prev_idx)
        except Exception:
            pass
        idx = int(sel[0])
        self.current_index = idx
        self.load_controls_from_settings(idx)
        self.render_preview(idx)
    def on_controls_changed(self):
        idx = self.current_index
        self.save_controls_to_settings(idx)
        self.render_preview(idx)
    def render_preview(self, idx: int):
        st = self.get_state()
        it = st["items"][idx]
        txtp = it["txt_path"]
        meta = it.get("meta", {}) or {}
        settings = st["settings"][idx]
        y_mode = settings["y_mode"]
        y_step = settings["y_step_kv"] if y_mode == "manual" else None
        df = read_measurement_txt(txtp)
        draw_oscillogram_on_axes(
            self.ax,
            df,
            st["serial_number"],
            st["divider_str"],
            meta,
            file_prefix=it.get("prefix"),
            divider1_str=st.get("divider1_str", ""),
            divider2_str=st.get("divider2_str", ""),
            show_matlab_header=True,
            textbox_mode_c=st.get("textbox_mode_c", "Auto (from Excel)"),
            xmax_us=settings["xmax_us"],
            y_mode=y_mode,
            y_step_kv=y_step,
            grid_scale=st["grid_scale"],
            grid_major_lw=float(parse_float_any(st.get("grid_major_lw", 0.85), 0.85)),
            grid_minor_lw=float(parse_float_any(st.get("grid_minor_lw", 0.0), 0.0)),
            plot_frame_lw=float(parse_float_any(st.get("plot_frame_lw", 1.2), 1.2)),
            textbox_frame_lw=float(parse_float_any(st.get("textbox_frame_lw", 1.0), 1.0)),
            tick_major_lw=float(parse_float_any(st.get("tick_major_lw", 1.0), 1.0)),
            tick_minor_lw=float(parse_float_any(st.get("tick_minor_lw", 0.8), 0.8)),
            font_label=int(parse_float_any(st.get("font_label", 6), 6)),
            font_tick=int(parse_float_any(st.get("font_tick", 6), 6)),
            font_header=int(parse_float_any(st.get("font_header", 6), 6)),
            font_textbox=int(parse_float_any(st.get("font_textbox", 6), 6)),
            textbox_pos_mode=str(st.get("textbox_pos_mode", "auto")),
            y_scale=float(parse_float_any(settings.get("scale_factor", 1.0), 1.0)),
        )
        m_left = float(parse_float_any(st.get("margin_left", 0.12), 0.12))
        m_right = float(parse_float_any(st.get("margin_right", 0.93), 0.93))
        m_bottom = float(parse_float_any(st.get("margin_bottom", 0.22), 0.22))
        m_top = float(parse_float_any(st.get("margin_top", 0.88), 0.88))
        self.fig.subplots_adjust(left=m_left, right=m_right, bottom=m_bottom, top=m_top)
        self.canvas.draw()
    def _current_label_suggest(self) -> str:
        st = self.get_state()
        try:
            it = st["items"][self.current_index]
            meta = it.get("meta", {}) or {}
            no = str(meta.get("No", it.get("No", ""))).strip()
            shape = normalize_shape(str(meta.get("Shape", "")))
            pref = str(it.get("prefix", "")).strip()
            parts = [p for p in [pref + no if pref and no else no, shape] if p]
            base = "_".join(parts) if parts else f"osc_{self.current_index+1}"
            base = re.sub(r"[^A-Za-z0-9_\-]+", "_", base)
            return base or f"osc_{self.current_index+1}"
        except Exception:
            return f"osc_{self.current_index+1}"
    def save_current_as_jpg(self):
        # Ensure current settings are applied
        try:
            self.on_controls_changed()
        except Exception:
            pass
        default_name = self._current_label_suggest() + ".jpg"
        out_path = filedialog.asksaveasfilename(
            title="Save oscillogram as JPG",
            defaultextension=".jpg",
            initialfile=default_name,
            filetypes=[("JPEG image", "*.jpg;*.jpeg"), ("All files", "*.*")],
        )
        if not out_path:
            return
        try:
            save_figure_as_jpg(self.fig, out_path, dpi=int(self.fig.dpi))
            messagebox.showinfo("Saved", f"Saved JPG:\n{out_path}")
        except Exception as e:
            messagebox.showerror(
                "Save failed",
                "Could not save JPG.\n\n"
                "Tip: install Pillow (required for JPG in many setups):\n"
                "  py -m pip install pillow\n\n"
                f"Details: {e}",
            )
    def copy_current_to_clipboard(self):
        # Ensure current settings are applied
        try:
            self.on_controls_changed()
        except Exception:
            pass
        try:
            copy_figure_to_clipboard_windows(self.fig, dpi=int(self.fig.dpi))
            messagebox.showinfo("Copied", "Oscillogram copied to clipboard.\nYou can paste it into Word/PowerPoint (Ctrl+V).")
        except Exception as e:
            messagebox.showerror(
                "Copy failed",
                "Could not copy image to clipboard.\n\n"
                "This feature requires Windows clipboard support packages:\n"
                "  py -m pip install pillow pywin32\n\n"
                f"Details: {e}",
            )
    def refresh_current(self):
        if not self.get_state().get("items"):
            return
        self.on_controls_changed()
    def apply_current_to_all(self):
        st = self.get_state()
        idx = self.current_index
        self.save_controls_to_settings(idx)
        cur = st["settings"][idx]
        for i in range(len(st["settings"])):
            st["settings"][i]["xmax_us"] = cur["xmax_us"]
            st["settings"][i]["y_mode"] = cur["y_mode"]
            st["settings"][i]["y_step_kv"] = cur["y_step_kv"]
            st["settings"][i]["scale_factor"] = cur.get("scale_factor", 1.0)
        messagebox.showinfo("Applied", "Current settings applied to ALL oscillograms.")
    def reset_selected(self):
        st = self.get_state()
        i = self.current_index
        st["settings"][i] = {
            "xmax_us": (st.get("default_xmax_c") if st.get("prefixes", [])[i] == "c" else st["default_xmax"]),
            "y_mode": "auto",
            "y_step_kv": None,
            "scale_factor": 1.0,
        }
        self.load_controls_from_settings(i)
        self.render_preview(i)
# =========================
# Main App V3.6
# =========================
class AppV3_3:
    def __init__(self, root: tk.Tk):
        self.root = root
        # Basic/Pro edition gate
        self._pro_unlocked = tk.BooleanVar(value=False)
        self._pro_widgets: list[tuple[tk.Widget, str]] = []  # (widget, unlocked_state)
        self._basic_keep_widgets: list[tk.Widget] = []  # widgets that stay enabled in BASIC
        self._edition_label: tk.Widget | None = None
        root.title("Oscillogram Generator (V3.7.31)")
        self.excel_path = tk.StringVar()
        self.template_path = tk.StringVar()
        # Auto-load template from same folder as .py (TEMPLATE.docx or first .docx)
        _auto_tpl = auto_find_template_docx()
        if _auto_tpl:
            self.template_path.set(_auto_tpl)
        self.serial_number_var = tk.StringVar()  # optional override
        # Legacy divider used by the existing top-right header (kept as requested)
        self.divider_str = tk.StringVar(value="3394,3")
        # NEW: optional folder source for v*/c* TXT files
        self.txt_folder = tk.StringVar()
        # NEW: MATLAB-style headers require Divider1/Divider2
        self.divider1_str = tk.StringVar(value="3394,3")
        self.divider2_str = tk.StringVar(value="3394,3")
        # Editable divider options (PRO can add/remove; persisted)
        self.divider_values: list[str] = ["3394,3", "864,0", "166,81", "458,24", "1117,7", "152,97", "259,36", "25,24", "49,28", "7409,13", "3704,25", "283,9", "2500"]
        self.default_xmax = tk.StringVar(value="100")
        self.default_xmax_c = tk.StringVar(value="100")  # default for c*.txt
        self.textbox_mode_c = tk.StringVar(value="Auto (from Excel)")  # textbox template for c*.txt
        self.default_y_step = tk.StringVar(value="50")
        # fixed size by request
        self.image_width_cm = tk.StringVar(value="9")
        self.image_height_cm = tk.StringVar(value="5,39")
        self.lock_aspect = tk.BooleanVar(value=False)
        # grid density
        self.grid_scale = tk.StringVar(value="Normal")  # Normal / Fine / Coarse
        # NEW: output DPI (sharpness)
        self.output_dpi = tk.StringVar(value="600")  # try 600; if heavy, use 300
        # NEW: export layout controls (MATLAB-like margins by default)
        # Values are in figure-relative coordinates [0..1]
        self.margin_left = tk.StringVar(value="0,12")
        self.margin_right = tk.StringVar(value="0,93")
        self.margin_bottom = tk.StringVar(value="0,22")
        self.margin_top = tk.StringVar(value="0,88")
        self.use_tight_bbox = tk.BooleanVar(value=False)
        self.pad_inches = tk.StringVar(value="0,02")
        
        # NEW: font tuning
        # NEW: grid linewidth tuning
        self.grid_major_lw = tk.StringVar(value="0,4")
        self.grid_minor_lw = tk.StringVar(value="0,0")
        # Frame linewidths
        self.plot_frame_lw = tk.StringVar(value="0,7")
        self.textbox_frame_lw = tk.StringVar(value="0,6")
        # Tick linewidths
        self.tick_major_lw = tk.StringVar(value="0,6")
        self.tick_minor_lw = tk.StringVar(value="0,4")
        # NEW: font tuning (points) for export/preview
        self.font_label = tk.StringVar(value="6")
        self.font_tick = tk.StringVar(value="6")
        self.font_header = tk.StringVar(value="6")
        self.font_textbox = tk.StringVar(value="6")
        # NEW: textbox placement
        self.textbox_pos_mode = tk.StringVar(value="auto")  # auto / bottom-right / top-right
        self.txt_paths: list[str] = []  # manually selected TXT files
        self.txt_index: dict[str, str] = {}  # basename(lower) -> full path
        self.items: list[dict] = []  # resolved work items (only existing files)
        self.meta_df: pd.DataFrame | None = None
        self.serial_number: str | None = None
        self.settings: list[dict] = []
        self._default_xmax_val = 100.0
        self.nb = ttk.Notebook(root)
        self.nb.pack(fill="both", expand=True)
        self.tab_input = ttk.Frame(self.nb)
        self.tab_osc = ttk.Frame(self.nb)
        self.tab_report = ttk.Frame(self.nb)
        self.nb.add(self.tab_input, text="1) Input")
        self.nb.add(self.tab_osc, text="2) Oscillograms")
        self.nb.add(self.tab_report, text="3) Report")
        self.build_tab_input()
        self.preview_panel = OscPreviewPanel(self.tab_osc, self.get_preview_state)
        self.preview_panel.pack(fill="both", expand=True)
        self.build_tab_report()
        # Apply BASIC lock initially
        self._apply_pro_lock()
        # Load persisted report/export defaults (Tab 3)
        self._defaults_path = self._get_defaults_path()
        self.load_report_defaults()
        # Persist settings on close
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    # -------------------------------
    # Basic / Pro gating
    # -------------------------------
    def _register_pro(self, widget: tk.Widget, unlocked_state: str = "normal") -> None:
        """Register a widget that is disabled in BASIC and enabled in PRO."""
        self._pro_widgets.append((widget, unlocked_state))

    def _set_widget_state_safe(self, widget: tk.Widget, state: str) -> None:
        try:
            widget.configure(state=state)
        except Exception:
            pass

    def _apply_pro_lock(self) -> None:
        """Apply current BASIC/PRO state to widgets."""
        is_pro = bool(self._pro_unlocked.get())

        # First, disable all registered PRO widgets in BASIC; enable in PRO.
        for w, unlocked_state in getattr(self, "_pro_widgets", []):
            if not w:
                continue
            if is_pro:
                self._set_widget_state_safe(w, unlocked_state)
            else:
                self._set_widget_state_safe(w, "disabled")

        # BASIC widgets always enabled
        for w in getattr(self, "_basic_keep_widgets", []):
            if not w:
                continue
            # try to restore sensible state
            # entries/buttons/checkbuttons can use "normal"; combobox often "readonly"
            st = "normal"
            try:
                if isinstance(w, ttk.Combobox):
                    st = "readonly"
            except Exception:
                st = "normal"
            self._set_widget_state_safe(w, st)

        # Update label if present
        try:
            if self._edition_label is not None:
                txt = "Edition: PRO" if is_pro else "Edition: BASIC"
                self._edition_label.configure(text=txt)
        except Exception:
            pass

    def unlock_pro(self) -> None:
        """Ask for password and unlock PRO features."""
        if bool(self._pro_unlocked.get()):
            messagebox.showinfo("PRO", "PRO is already unlocked.")
            return
        pw = simpledialog.askstring("Unlock PRO", "Unesi lozinku za PRO:", show="*")
        if pw is None:
            return
        # Simple check (can be overridden via env var OSCGEN_PRO_PASSWORD)
        if pw == OSCGEN_PRO_PASSWORD:
            self._pro_unlocked.set(True)
            self._apply_pro_lock()
            messagebox.showinfo("PRO", "PRO features unlocked.")
        else:
            messagebox.showerror("PRO", "Pogrešna lozinka.")

    # -------------------------------
    # Persisted defaults (Tab 3)
    # -------------------------------
    def _get_defaults_path(self) -> str:
        """Path for storing Tab 3 (Report) defaults JSON.
        We store next to the script so users can carry the defaults
        with the tool folder (e.g. OneDrive).
        """
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        except Exception:
            base_dir = os.getcwd()
        return os.path.join(base_dir, "oscilogram_generator_report_defaults.json")
    def _collect_report_defaults(self) -> dict:
        """Collect current Tab 3 knobs into a JSON-serializable dict."""
        # Note: keep keys stable; this is user-facing persistence.
        return {
            # size / dpi
            "image_width_cm": self.image_width_cm.get(),
            "image_height_cm": self.image_height_cm.get(),
            "output_dpi": self.output_dpi.get(),
            # margins
            "margin_left": self.margin_left.get(),
            "margin_right": self.margin_right.get(),
            "margin_bottom": self.margin_bottom.get(),
            "margin_top": self.margin_top.get(),
            "use_tight_bbox": bool(self.use_tight_bbox.get()),
            "pad_inches": self.pad_inches.get(),
            # fonts
            "font_label": self.font_label.get(),
            "font_tick": self.font_tick.get(),
            "font_header": self.font_header.get(),
            "font_textbox": self.font_textbox.get(),
            # grid
            "grid_major_lw": self.grid_major_lw.get(),
            "grid_minor_lw": self.grid_minor_lw.get(),
            # frames
            "plot_frame_lw": self.plot_frame_lw.get(),
            "textbox_frame_lw": self.textbox_frame_lw.get(),
            # ticks
            "tick_major_lw": self.tick_major_lw.get(),
            "tick_minor_lw": self.tick_minor_lw.get(),
            # textbox placement
            "textbox_pos_mode": self.textbox_pos_mode.get(),
            # divider dropdown options + current selections
            "divider_values": list(getattr(self, "divider_values", [])),
            "divider1_str": self.divider1_str.get(),
            "divider2_str": self.divider2_str.get(),
        }
    def _apply_report_defaults(self, d: dict) -> None:
        """Apply persisted defaults (best-effort)."""
        def _set(var: tk.Variable, key: str):
            if key in d and d[key] is not None:
                try:
                    var.set(d[key])
                except Exception:
                    pass
        _set(self.image_width_cm, "image_width_cm")
        _set(self.image_height_cm, "image_height_cm")
        _set(self.output_dpi, "output_dpi")
        _set(self.margin_left, "margin_left")
        _set(self.margin_right, "margin_right")
        _set(self.margin_bottom, "margin_bottom")
        _set(self.margin_top, "margin_top")
        if "use_tight_bbox" in d:
            try:
                self.use_tight_bbox.set(bool(d["use_tight_bbox"]))
            except Exception:
                pass
        _set(self.pad_inches, "pad_inches")
        _set(self.font_label, "font_label")
        _set(self.font_tick, "font_tick")
        _set(self.font_header, "font_header")
        _set(self.font_textbox, "font_textbox")
        _set(self.grid_major_lw, "grid_major_lw")
        _set(self.grid_minor_lw, "grid_minor_lw")
        _set(self.plot_frame_lw, "plot_frame_lw")
        _set(self.textbox_frame_lw, "textbox_frame_lw")
        _set(self.tick_major_lw, "tick_major_lw")
        _set(self.tick_minor_lw, "tick_minor_lw")
        # divider dropdown options + current selections
        try:
            if isinstance(d.get("divider_values"), list) and d.get("divider_values"):
                self.divider_values = [str(x).strip() for x in d.get("divider_values") if str(x).strip()]
        except Exception:
            pass
        _set(self.divider1_str, "divider1_str")
        _set(self.divider2_str, "divider2_str")
        try:
            self._refresh_divider_comboboxes()
        except Exception:
            pass
        _set(self.textbox_pos_mode, "textbox_pos_mode")
    def load_report_defaults(self) -> None:
        """Load defaults from JSON file (if it exists)."""
        path = getattr(self, "_defaults_path", None) or self._get_defaults_path()
        if not os.path.exists(path):
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                self._apply_report_defaults(data)
        except Exception:
            # If file is corrupt, ignore silently (user can delete it).
            return
    def save_report_defaults(self) -> None:
        """Save current defaults to JSON file."""
        path = getattr(self, "_defaults_path", None) or self._get_defaults_path()
        try:
            data = self._collect_report_defaults()
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception:
            return
    def on_close(self) -> None:
        """Save defaults and close the GUI."""
        try:
            self.save_report_defaults()
        finally:
            try:
                self.root.destroy()
            except Exception:
                pass
    def build_tab_input(self):
        f = self.tab_input
        row = 0
        ttk.Label(f, text="Metadata file (.xlsx or .csv):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(f, textvariable=self.excel_path, width=70).grid(row=row, column=1, padx=8, pady=6)
        ttk.Button(f, text="Browse", command=self.pick_excel).grid(row=row, column=2, padx=8, pady=6)
        row += 1
        ttk.Label(f, text="TXT files:").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        ttk.Button(f, text="Select TXT files", command=self.pick_txts).grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self.txt_label = ttk.Label(f, text="0 selected")
        self.txt_label.grid(row=row, column=2, sticky="w", padx=8, pady=6)
        row += 1
        ttk.Label(f, text="TXT source folder (optional):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        txt_folder_entry = ttk.Entry(f, textvariable=self.txt_folder, width=70)
        txt_folder_entry.grid(row=row, column=1, padx=8, pady=6)
        txt_folder_btn = ttk.Button(f, text="Browse", command=self.pick_txt_folder)
        txt_folder_btn.grid(row=row, column=2, padx=8, pady=6)
        row += 1
        ttk.Label(f, text="Divider1 (v header):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        div1_entry = ttk.Combobox(
            f,
            textvariable=self.divider1_str,
            state="readonly",
            width=18,
            values=self.divider_values
        )
        div1_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._div1_entry = div1_entry
        row += 1
        ttk.Label(f, text="Divider2 (c header):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        div2_entry = ttk.Combobox(
            f,
            textvariable=self.divider2_str,
            state="readonly",
            width=18,
            values=self.divider_values
        )
        div2_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._div2_entry = div2_entry
        row += 1

        # PRO: allow editing divider dropdown options (persisted)
        divider_btns = ttk.Frame(f)
        divider_btns.grid(row=row, column=1, sticky="w", padx=8, pady=(0, 6))
        add_div_btn = ttk.Button(divider_btns, text="Add divider value…", command=self.add_divider_value)
        add_div_btn.pack(side="left")
        rm_div_btn = ttk.Button(divider_btns, text="Remove divider value…", command=self.remove_divider_value)
        rm_div_btn.pack(side="left", padx=8)
        self._register_pro(add_div_btn, "normal")
        self._register_pro(rm_div_btn, "normal")
        row += 1
        ttk.Label(f, text="Word template (.docx):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(f, textvariable=self.template_path, width=70).grid(row=row, column=1, padx=8, pady=6)
        ttk.Button(f, text="Browse", command=self.pick_template).grid(row=row, column=2, padx=8, pady=6)
        row += 1
        ttk.Label(f, text="Serial number (tested):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(f, textvariable=self.serial_number_var, width=30).grid(row=row, column=1, sticky="w", padx=8, pady=6)
        ttk.Label(f, text="(optional; default = Excel filename)").grid(row=row, column=2, sticky="w", padx=8, pady=6)
        row += 1
        ttk.Separator(f).grid(row=row, column=0, columnspan=3, sticky="ew", padx=8, pady=10)
        row += 1
        ttk.Label(f, text="Default X max (us) [v]:").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        xmaxv_entry = ttk.Entry(f, textvariable=self.default_xmax, width=10)
        xmaxv_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(xmaxv_entry, "normal")
        row += 1

        ttk.Label(f, text="Default X max (us) [c]:").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        xmaxc_entry = ttk.Entry(f, textvariable=self.default_xmax_c, width=10)
        xmaxc_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(xmaxc_entry, 'normal')
        row += 1


        ttk.Label(f, text="Textbox template [c]:").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        cb = ttk.Combobox(
            f,
            textvariable=self.textbox_mode_c,
            state="readonly",
            width=70,
            values=[
                "Auto (from Excel)",
                "Default",
                "LI full (Upk,T1,T2)",
                "LI full (Ipk)",
                "LI full (Ipk max/min)",
                "LI tailchopped (Upk min/max, T1, T2--, Tc)",
                "LI tailchopped (Ipk)",
                "SI IEC 60060 (Upk,Tp,T2)",
                "SI IEC 60060 (Ipk,Tp,T2)",
                "SI IEC 60076 (Upk,T1,Td,T0)",
                "SI IEC 60076 (Ipk max/min)",
            ],
        )
        cb.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(cb, 'readonly')
        row += 1

        ttk.Label(f, text="Default Manual Y step (kV):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        ystep_entry = ttk.Entry(f, textvariable=self.default_y_step, width=10)
        ystep_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(ystep_entry, "normal")
        row += 1
        ttk.Label(f, text="Image width (cm):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        w_entry = ttk.Entry(f, textvariable=self.image_width_cm, width=10)
        w_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(w_entry, "normal")
        row += 1
        ttk.Label(f, text="Image height (cm):").grid(row=row, column=0, sticky="w", padx=8, pady=6)
        h_entry = ttk.Entry(f, textvariable=self.image_height_cm, width=10)
        h_entry.grid(row=row, column=1, sticky="w", padx=8, pady=6)
        self._register_pro(h_entry, "normal")
        row += 1
        lock_cb = ttk.Checkbutton(f, text="Lock aspect ratio", variable=self.lock_aspect)
        lock_cb.grid(row=row, column=1, sticky="w", padx=8, pady=(0, 8))
        self._register_pro(lock_cb, "normal")
        row += 1
        ttk.Button(f, text="Load inputs (prepare preview)", command=self.load_everything).grid(
            row=row, column=1, sticky="w", padx=8, pady=14
        )

        # -------------------------------------------------
        # BASIC / PRO status + unlock button
        # -------------------------------------------------
        row += 1
        ttk.Separator(f).grid(row=row, column=0, columnspan=3, sticky="ew", padx=8, pady=12)
        row += 1

        self._edition_label = ttk.Label(f, text="Edition: BASIC", font=("Segoe UI", 9, "bold"))
        self._edition_label.grid(row=row, column=0, sticky="w", padx=8, pady=6)

        unlock_btn = ttk.Button(f, text="Unlock PRO", command=self.unlock_pro)
        unlock_btn.grid(row=row, column=1, sticky="w", padx=8, pady=6)

        # Gumb mora uvijek ostati aktivan
        self._basic_keep_widgets.append(unlock_btn)

        f.grid_columnconfigure(1, weight=1)
    def build_tab_report(self):
        f = self.tab_report
        ttk.Label(
            f,
            text="Export tuning (these settings affect Word/JPG/Clipboard export look)",
        ).grid(row=0, column=0, columnspan=4, sticky="w", padx=10, pady=(12, 8))
        r = 1
        ttk.Label(f, text="Margins (0..1):").grid(row=r, column=0, sticky="w", padx=10, pady=(12, 6))
        r += 1
        ttk.Label(f, text="Left:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.margin_left, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(f, text="Right:").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.margin_right, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Bottom:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.margin_bottom, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(f, text="Top:").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.margin_top, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Grid linewidth:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Label(f, text="Major:").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.grid_major_lw, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Minor grid linewidth (0 = off):").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.grid_minor_lw, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Frame linewidth:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Label(f, text="Plot: ").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.plot_frame_lw, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Textbox border linewidth:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.textbox_frame_lw, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Tick linewidth:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Label(f, text="Major: ").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.tick_major_lw, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Minor tick linewidth:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.tick_minor_lw, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Checkbutton(f, text="Tight crop (bbox_inches='tight')", variable=self.use_tight_bbox).grid(
            row=r, column=0, columnspan=2, sticky="w", padx=10, pady=(12, 6)
        )
        ttk.Label(f, text="Pad (in):").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=(12, 6))
        ttk.Entry(f, textvariable=self.pad_inches, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=(12, 6))
        
        r += 1
        ttk.Label(f, text="Fonts (pt):").grid(row=r, column=0, sticky="w", padx=10, pady=(12, 6))
        r += 1
        ttk.Label(f, text="Label:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.font_label, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(f, text="Tick:").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.font_tick, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Header:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        ttk.Entry(f, textvariable=self.font_header, width=10).grid(row=r, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(f, text="Textbox:").grid(row=r, column=2, sticky="w", padx=(18, 6), pady=6)
        ttk.Entry(f, textvariable=self.font_textbox, width=10).grid(row=r, column=3, sticky="w", padx=6, pady=6)
        r += 1
        ttk.Label(f, text="Textbox position:").grid(row=r, column=0, sticky="w", padx=10, pady=6)
        cb = ttk.Combobox(f, textvariable=self.textbox_pos_mode, width=18, state="readonly",
                         values=("auto", "bottom-right", "top-right"))
        cb.grid(row=r, column=1, sticky="w", padx=6, pady=6)
        ttk.Label(f, text="(auto = based on polarity)").grid(row=r, column=2, columnspan=2, sticky="w", padx=(18, 6), pady=6)
        r += 1
        btnrow = ttk.Frame(f)
        btnrow.grid(row=r, column=0, columnspan=4, sticky="w", padx=10, pady=(12, 6))
        ttk.Button(btnrow, text="Export TEST JPG (selected)…", command=self.export_test_jpg_selected).pack(side="left")
        ttk.Button(btnrow, text="Apply export settings to Preview", command=self.apply_export_settings_to_preview).pack(side="left", padx=10)
        r += 1
        ttk.Button(f, text="Generate Word (Save As...)", command=self.generate_word).grid(
            row=r, column=0, columnspan=4, sticky="w", padx=10, pady=(14, 8)
        )
        r += 1

        # --- BASIC/PRO gating: everything in Tab 3 except size + lock aspect + Generate Word is PRO ---
        # We register widgets here so they can be disabled in BASIC.
        try:
            keep = set(self._basic_keep_widgets)
            def walk(parent):
                for ch in parent.winfo_children():
                    # Keep only specific basics:
                    if ch in keep:
                        pass
                    else:
                        # Keep the main "Generate Word" button enabled in BASIC
                        try:
                            if isinstance(ch, ttk.Button) and (ch.cget("text") or "").startswith("Generate Word"):
                                self._basic_keep_widgets.append(ch)
                            else:
                                # anything else interactive becomes PRO
                                if isinstance(ch, (ttk.Entry, ttk.Combobox, ttk.Checkbutton, ttk.Button, ttk.Radiobutton)):
                                    # combobox should restore to readonly
                                    st = "readonly" if isinstance(ch, ttk.Combobox) else "normal"
                                    self._register_pro(ch, st)
                        except Exception:
                            pass
                    walk(ch)
            walk(f)
        except Exception:
            pass

        ttk.Label(
            f,
            text=(
                "Tips:\n"
                "• If it looks blurry in Word → increase DPI (e.g. 600).\n"
                "• If plot area is too wide/narrow → tweak Left/Right margins.\n"
                "• Tight crop removes whitespace automatically (may change the MATLAB-like look)."
            ),
        ).grid(row=r, column=0, columnspan=4, sticky="w", padx=10, pady=(4, 10))
        for c in range(4):
            f.grid_columnconfigure(c, weight=1)
    def _read_export_tuning(self):
        """Read export tuning values from UI (Tab 3)."""
        return dict(
            width_cm=parse_float_any(self.image_width_cm.get(), 9.61),
            height_cm=parse_float_any(self.image_height_cm.get(), 5.37),
            lock_aspect=bool(self.lock_aspect.get()),
            grid_scale=self.grid_scale.get(),
            out_dpi=int(parse_float_any(self.output_dpi.get(), 600)),
            margin_left=parse_float_any(self.margin_left.get(), 0.12),
            margin_right=parse_float_any(self.margin_right.get(), 0.93),
            margin_bottom=parse_float_any(self.margin_bottom.get(), 0.22),
            margin_top=parse_float_any(self.margin_top.get(), 0.88),
            grid_major_lw=parse_float_any(self.grid_major_lw.get(), 0.85),
            grid_minor_lw=parse_float_any(self.grid_minor_lw.get(), 0.0),
            plot_frame_lw=parse_float_any(self.plot_frame_lw.get(), 1.2),
            textbox_frame_lw=parse_float_any(self.textbox_frame_lw.get(), 1.0),
            tick_major_lw=parse_float_any(self.tick_major_lw.get(), 1.0),
            tick_minor_lw=parse_float_any(self.tick_minor_lw.get(), 0.8),
            use_tight_bbox=bool(self.use_tight_bbox.get()),
            pad_inches=parse_float_any(self.pad_inches.get(), 0.02),
            font_label=int(parse_float_any(self.font_label.get(), 6)),
            font_tick=int(parse_float_any(self.font_tick.get(), 6)),
            font_header=int(parse_float_any(self.font_header.get(), 6)),
            font_textbox=int(parse_float_any(self.font_textbox.get(), 6)),
            textbox_pos_mode=str(self.textbox_pos_mode.get()),
        )
    def apply_export_settings_to_preview(self):
        """Apply current Tab 3 export tuning to the Preview (Tab 2) and refresh."""
        try:
            if hasattr(self, "preview_panel") and self.preview_panel is not None:
                tuning = self._read_export_tuning()
                st = self.preview_panel.get_state()
                # only propagate visual tuning (not sizes) to the preview renderer
                for k in [
                    "grid_scale",
                    "grid_major_lw",
                    "grid_minor_lw",
                    "plot_frame_lw",
                    "textbox_frame_lw",
                    "tick_major_lw",
                    "tick_minor_lw",
                    "margin_left", "margin_right", "margin_bottom", "margin_top",
                    "font_label", "font_tick", "font_header", "font_textbox",
                ]:
                    if k in tuning:
                        st[k] = tuning[k]
                self.preview_panel.refresh_current()
        except Exception as e:
            messagebox.showerror("Preview refresh failed", f"Could not refresh preview.\n\nDetails: {e}")
    def export_test_jpg_selected(self):
        """Export a single TEST JPG for the currently selected oscillogram using Tab 3 settings."""
        try:
            self.ensure_loaded()
            if not hasattr(self, "preview_panel") or self.preview_panel is None:
                messagebox.showerror("No preview", "Preview panel not available.")
                return
            idx = getattr(self.preview_panel, "current_index", 0)
            if idx < 0 or idx >= len(self.items):
                messagebox.showerror("No selection", "Please select an oscillogram in Tab 2 first.")
                return
            it = self.items[idx]
            meta = it.get("meta", {}) or {}
            st = self.settings[idx]
            y_mode = st["y_mode"]
            y_step = st["y_step_kv"] if (y_mode == "manual") else None
            tuning = self._read_export_tuning()
            default_name = f"TEST_{self.serial_number}_{it.get('prefix','x')}{it.get('No','')}.jpg"
            out_path = filedialog.asksaveasfilename(
                title="Export TEST JPG",
                defaultextension=".jpg",
                initialfile=default_name,
                filetypes=[("JPEG image", "*.jpg;*.jpeg"), ("All files", "*.*")],
            )
            if not out_path:
                return
            tmpdir = tempfile.mkdtemp(prefix="osc_test_")
            tmp_png = os.path.join(tmpdir, "tmp.png")
            make_oscillogram_png(
                txt_path=it["txt_path"],
                out_png=tmp_png,
                serial_number=self.serial_number,
                divider_str=self.divider_str.get().strip(),
                meta=meta,
                file_prefix=it.get("prefix"),
                divider1_str=self.divider1_str.get().strip(),
                divider2_str=self.divider2_str.get().strip(),
                show_matlab_header=True,
                textbox_mode_c=self.textbox_mode_c.get(),
                xmax_us=float(st["xmax_us"]),
                y_mode=y_mode,
                y_step_kv=y_step,
                grid_scale=tuning["grid_scale"],
                grid_major_lw=tuning.get("grid_major_lw", 0.85),
                grid_minor_lw=tuning.get("grid_minor_lw", 0.0),
                plot_frame_lw=tuning.get("plot_frame_lw", 1.2),
                textbox_frame_lw=tuning.get("textbox_frame_lw", 1.0),
                tick_major_lw=tuning.get("tick_major_lw", 1.0),
                tick_minor_lw=tuning.get("tick_minor_lw", 0.8),
                out_width_cm=tuning["width_cm"],
                out_height_cm=tuning["height_cm"],
                out_dpi=tuning["out_dpi"],
                font_label=tuning["font_label"],
                font_tick=tuning["font_tick"],
                font_header=tuning["font_header"],
                font_textbox=tuning["font_textbox"],
                textbox_pos_mode=tuning.get("textbox_pos_mode", "auto"),
                margin_left=tuning["margin_left"],
                margin_right=tuning["margin_right"],
                margin_bottom=tuning["margin_bottom"],
                margin_top=tuning["margin_top"],
                use_tight_bbox=tuning["use_tight_bbox"],
                pad_inches=tuning["pad_inches"],
            )
            try:
                from PIL import Image  # type: ignore
                img = Image.open(tmp_png).convert("RGB")
                _ensure_parent_dir(out_path)
                img.save(out_path, "JPEG", quality=95)
            except Exception as e:
                messagebox.showerror(
                    "Export failed",
                    "Could not export TEST JPG.\n\n"
                    "Tip: install Pillow:\n  py -m pip install pillow\n\n"
                    f"Details: {e}",
                )
                return
            messagebox.showinfo("Exported", f"Saved TEST JPG:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Export failed", f"Could not export TEST JPG.\n\nDetails: {e}")
    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[
            ("Excel or CSV", "*.xlsx;*.csv"),
            ("Excel files", "*.xlsx"),
            ("CSV files", "*.csv"),
            ("All files", "*.*"),
        ])
        if p:
            self.excel_path.set(p)
    def pick_template(self):
        base = str(get_base_dir())
        p = filedialog.askopenfilename(
            initialdir=base,
            filetypes=[("Word files", "*.docx")]
        )
        if p:
            self.template_path.set(p)
    def pick_txts(self):
        ps = filedialog.askopenfilenames(filetypes=[("TXT files", "*.txt;*.TXT")])
        self.txt_paths = list(ps)
        self.txt_label.config(text=f"{len(self.txt_paths)} selected")
    def pick_txt_folder(self):
        p = filedialog.askdirectory(title="Select folder with v*/c* TXT files")
        if p:
            self.txt_folder.set(p)


    # ---------- divider options (PRO) ----------
    def _normalize_divider_value(self, s: str) -> str:
        """Normalize divider value input (keep decimal comma style)."""
        s = (s or "").strip()
        if not s:
            return ""
        # allow users to type '.' or ','
        s = s.replace(" ", "")
        # keep only the numeric-looking part
        m = re.search(r"[-+]?\d+(?:[\.,]\d+)?", s)
        if not m:
            return ""
        val = m.group(0).replace(".", ",")
        # strip trailing comma
        val = val.rstrip(",")
        return val

    def _divider_value_key(self, s: str) -> float:
        """Numeric key for sorting divider values (best-effort)."""
        try:
            return float(str(s).replace(",", "."))
        except Exception:
            return float("inf")

    def _refresh_divider_comboboxes(self) -> None:
        """Apply current divider_values to both dropdowns."""
        vals = list(self.divider_values or [])
        try:
            if hasattr(self, "_div1_entry") and self._div1_entry is not None:
                self._div1_entry.configure(values=vals)
        except Exception:
            pass
        try:
            if hasattr(self, "_div2_entry") and self._div2_entry is not None:
                self._div2_entry.configure(values=vals)
        except Exception:
            pass

    def add_divider_value(self) -> None:
        """PRO: Add a new divider value to the dropdown list (persisted on close)."""
        s = simpledialog.askstring("Add divider value", "Upiši novu vrijednost djelila (npr. 3394,3):")
        if s is None:
            return
        v = self._normalize_divider_value(s)
        if not v:
            messagebox.showerror("Invalid", "Nisam prepoznao broj. Pokušaj npr. 3394,3")
            return
        if v in self.divider_values:
            messagebox.showinfo("Exists", f"Vrijednost {v} već postoji u listi.")
            return
        self.divider_values.append(v)
        # keep list sorted numerically for nicer UX
        try:
            self.divider_values = sorted(self.divider_values, key=self._divider_value_key)
        except Exception:
            pass
        self._refresh_divider_comboboxes()
        try:
            self.save_report_defaults()
        except Exception:
            pass
        messagebox.showinfo("Added", f"Dodano: {v}\n\nLista će se spremiti pri izlasku iz programa.")

    def remove_divider_value(self) -> None:
        """PRO: Remove a divider value from the dropdown list (persisted on close)."""
        if not self.divider_values:
            messagebox.showinfo("Empty", "Lista je prazna.")
            return
        cur_list = ", ".join(self.divider_values)
        s = simpledialog.askstring(
            "Remove divider value",
            "Upiši vrijednost koju želiš maknuti:\n\nTrenutno: " + cur_list,
        )
        if s is None:
            return
        v = self._normalize_divider_value(s)
        if not v or v not in self.divider_values:
            messagebox.showerror("Not found", f"Vrijednost '{s}' nije pronađena u listi.")
            return
        # prevent removing the currently selected values (auto-switch first)
        try:
            if self.divider1_str.get().strip() == v:
                # fallback to first remaining after removal
                pass
            if self.divider2_str.get().strip() == v:
                pass
        except Exception:
            pass
        self.divider_values = [x for x in self.divider_values if x != v]
        # if current selections were removed, set to a safe value
        if self.divider1_str.get().strip() == v:
            self.divider1_str.set(self.divider_values[0] if self.divider_values else "")
        if self.divider2_str.get().strip() == v:
            self.divider2_str.set(self.divider_values[0] if self.divider_values else "")
        self._refresh_divider_comboboxes()
        try:
            self.save_report_defaults()
        except Exception:
            pass
        messagebox.showinfo("Removed", f"Maknuto: {v}\n\nLista će se spremiti pri izlasku iz programa.")
    # ---------- loading ----------
    def ensure_loaded(self):
        excel = self.excel_path.get()
        if not excel or not Path(excel).exists():
            raise ValueError("Please select a valid metadata file (.xlsx or .csv).")
        ext = Path(excel).suffix.lower()
        if ext not in (".xlsx", ".csv"):
            raise ValueError("Metadata file must be .xlsx or .csv")
        template = self.template_path.get()
        if not template or not Path(template).exists():
            raise ValueError("Please choose a Word template (.docx).")
        # TXT sources: folder and/or manual selection
        folder = self.txt_folder.get().strip()
        has_folder = bool(folder) and Path(folder).exists() and Path(folder).is_dir()
        has_manual = bool(self.txt_paths)
        if not has_folder and not has_manual:
            raise ValueError("Please select TXT files and/or a TXT source folder.")
        divider = self.divider_str.get().strip()
        if not divider:
            raise ValueError("Please enter Divider value (legacy).")
        # Divider1/Divider2 (MATLAB headers)
        if not self.divider1_str.get().strip():
            self.divider1_str.set(divider)
        if not self.divider2_str.get().strip():
            self.divider2_str.set(divider)
        default_xmax = parse_float_any(self.default_xmax.get(), 100.0)
        default_xmax_c = parse_float_any(self.default_xmax_c.get(), default_xmax)
        default_y_step = parse_float_any(self.default_y_step.get(), 50.0)
        # Load Excel metadata (full set of columns used for textbox logic)
        self.meta_df = load_excel_metadata(excel)
        _sn = (self.serial_number_var.get() or "").strip()
        self.serial_number = _sn if _sn else Path(excel).stem
        # Build TXT index: manual files override folder; folder chooses newest on duplicates
        self.txt_index = {}
        # 1) folder scan (one level, no subfolders)
        folder_dupes: dict[str, list[str]] = {}
        if has_folder:
            for p in Path(folder).glob("*.txt"):
                key = p.name.lower()
                existing = self.txt_index.get(key)
                if existing:
                    folder_dupes.setdefault(key, [existing]).append(str(p))
                    try:
                        # keep newest
                        if p.stat().st_mtime > Path(existing).stat().st_mtime:
                            self.txt_index[key] = str(p)
                    except Exception:
                        pass
                else:
                    self.txt_index[key] = str(p)
        # 2) manual selection overrides
        for p in self.txt_paths:
            key = Path(p).name.lower()
            self.txt_index[key] = str(p)
        # Resolve work items based on Excel row index r (1-based) and No using criterion B
        self.items = []
        missing = []
        if self.meta_df is None:
            raise ValueError("Excel metadata could not be loaded.")
        for i in range(len(self.meta_df)):
            r = i + 1  # MATLAB-style 1-based row index
            meta = self.meta_df.iloc[i].to_dict()
            meta["No"] = str(meta.get("No", "")).strip()
            try:
                no = int(parse_float_any(meta.get("No"), 0))
            except Exception:
                no = 0
            if no <= 0:
                continue
            # Criterion B (MATLAB): decide v/c by position vs No
            if (r == no) or (r == (no * 2 - 1)):
                prefix = "v"
            else:
                prefix = "c"
            expected = f"{prefix}{no}.txt".lower()
            txt_path = self.txt_index.get(expected)
            if not txt_path or not Path(txt_path).exists():
                missing.append(expected)
                continue
            self.items.append({
                "excel_index": i,
                "r": r,
                "No": no,
                "prefix": prefix,
                "expected_name": expected,
                "txt_path": txt_path,
                "meta": meta,
            })
        if not self.items:
            raise ValueError(
                "No matching TXT files were found for the Excel rows.\n"
                "Tip: ensure files are named like v<No>.txt / c<No>.txt and are in the selected folder or manual list."
            )
        # Initialize per-item settings (xmax / y mode / y step)
        if not self.settings or len(self.settings) != len(self.items):
            self.settings = [{
                "xmax_us": (default_xmax_c if self.items[i].get("prefix") == "c" else default_xmax),
                "y_mode": "auto",
                "y_step_kv": default_y_step,
                "scale_factor": 1.0,
            } for i in range(len(self.items))]
        self._default_xmax_val = default_xmax
        self._default_xmax_c_val = default_xmax_c
    def load_everything(self):
        try:
            # Re-apply current defaults (xmax etc.) on every load
            self.settings = []
            self.ensure_loaded()
            self.preview_panel.load_list()
            self.nb.select(self.tab_osc)
            messagebox.showinfo("Loaded", "Inputs loaded. Preview is ready on Tab 2.")
        except Exception as e:
            messagebox.showerror("Error", str(e))
    def get_preview_state(self) -> dict:
        if self.meta_df is None or self.serial_number is None or not self.items:
            return {}
        return {
            "items": self.items,
            "meta_df": self.meta_df,
            "serial_number": self.serial_number,
            "divider_str": self.divider_str.get().strip(),  # legacy header divider (kept)
            "divider1_str": self.divider1_str.get().strip(),
            "divider2_str": self.divider2_str.get().strip(),
            "settings": self.settings,
            "default_xmax": self._default_xmax_val,
            "default_xmax_c": getattr(self, "_default_xmax_c_val", self._default_xmax_val),
            "textbox_mode_c": self.textbox_mode_c.get(),
            "prefixes": [it.get("prefix") for it in self.items],
            "grid_scale": self.grid_scale.get(),
            "grid_major_lw": self.grid_major_lw.get(),
            "grid_minor_lw": self.grid_minor_lw.get(),
            "plot_frame_lw": self.plot_frame_lw.get(),
            "textbox_frame_lw": self.textbox_frame_lw.get(),
            "tick_major_lw": self.tick_major_lw.get(),
            "tick_minor_lw": self.tick_minor_lw.get(),
            "margin_left": self.margin_left.get(),
            "margin_right": self.margin_right.get(),
            "margin_bottom": self.margin_bottom.get(),
            "margin_top": self.margin_top.get(),
            "font_label": self.font_label.get(),
            "font_tick": self.font_tick.get(),
            "font_header": self.font_header.get(),
            "font_textbox": self.font_textbox.get(),
            "textbox_pos_mode": self.textbox_pos_mode.get(),
            # divider dropdown options + current selections
            "divider_values": list(getattr(self, "divider_values", [])),
            "divider1_str": self.divider1_str.get(),
            "divider2_str": self.divider2_str.get(),
        }
    # ---------- report ----------
    def generate_word(self):
        try:
            self.ensure_loaded()
            template = self.template_path.get()
            # Persist current Tab 3 tuning so it becomes the default next time
            self.save_report_defaults()
            width_cm = parse_float_any(self.image_width_cm.get(), 9.61)
            height_cm = parse_float_any(self.image_height_cm.get(), 5.37)
            lock_aspect = bool(self.lock_aspect.get())
            gscale = self.grid_scale.get()
            out_dpi = int(parse_float_any(self.output_dpi.get(), 600))
            g_major_lw = parse_float_any(self.grid_major_lw.get(), 0.85)
            g_minor_lw = parse_float_any(self.grid_minor_lw.get(), 0.0)
            plot_frame_lw = parse_float_any(self.plot_frame_lw.get(), 1.2)
            textbox_frame_lw = parse_float_any(self.textbox_frame_lw.get(), 1.0)
            m_left = parse_float_any(self.margin_left.get(), 0.12)
            m_right = parse_float_any(self.margin_right.get(), 0.93)
            m_bottom = parse_float_any(self.margin_bottom.get(), 0.22)
            m_top = parse_float_any(self.margin_top.get(), 0.88)
            use_tight = bool(self.use_tight_bbox.get())
            pad_in = parse_float_any(self.pad_inches.get(), 0.02)
            f_label = int(parse_float_any(self.font_label.get(), 6))
            f_tick = int(parse_float_any(self.font_tick.get(), 6))
            f_header = int(parse_float_any(self.font_header.get(), 6))
            f_textbox = int(parse_float_any(self.font_textbox.get(), 6))
            # Tick widths
            t_major_lw = parse_float_any(self.tick_major_lw.get(), 1.0)
            t_minor_lw = parse_float_any(self.tick_minor_lw.get(), 0.8)
            # Textbox placement
            tb_pos_mode = str(self.textbox_pos_mode.get())
            suggested = f"OUTPUT_{self.serial_number}.docx"
            out_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=suggested,
                filetypes=[("Word document", "*.docx")]
            )
            if not out_path:
                return
            tmpdir = tempfile.mkdtemp(prefix="osc_v3_6_")
            images: list[str] = []
            peaks_table_rows: list[dict] = []
            for i, it in enumerate(self.items):
                meta = it.get("meta", {}) or {}
                st = self.settings[i]
                y_mode = st["y_mode"]
                y_step = st["y_step_kv"] if (y_mode == "manual") else None
                out_png = os.path.join(tmpdir, f"osc_{i+1:03d}.png")
                make_oscillogram_png(
                    txt_path=it["txt_path"],
                    out_png=out_png,
                    serial_number=self.serial_number,
                    divider_str=self.divider_str.get().strip(),
                    meta=meta,
                    file_prefix=it.get("prefix"),
                    divider1_str=self.divider1_str.get().strip(),
                    divider2_str=self.divider2_str.get().strip(),
                    show_matlab_header=True,
                    textbox_mode_c=self.textbox_mode_c.get(),
                    xmax_us=float(st["xmax_us"]),
                    y_mode=y_mode,
                    y_step_kv=y_step,
                    grid_scale=gscale,
                    grid_major_lw=g_major_lw,
                    grid_minor_lw=g_minor_lw,
                    plot_frame_lw=plot_frame_lw,
                    textbox_frame_lw=textbox_frame_lw,
                    out_width_cm=width_cm,
                    out_height_cm=height_cm,
                    out_dpi=out_dpi,
                    font_label=f_label,
                    font_tick=f_tick,
                    font_header=f_header,
                    font_textbox=f_textbox,
                    tick_major_lw=t_major_lw,
                    tick_minor_lw=t_minor_lw,
                    textbox_pos_mode=tb_pos_mode,
                    margin_left=m_left,
                    margin_right=m_right,
                    margin_bottom=m_bottom,
                    margin_top=m_top,
                    use_tight_bbox=use_tight,
                    pad_inches=pad_in,
                    y_scale=float(parse_float_any(st.get("scale_factor", 1.0), 1.0)),
                )
                # Collect textbox values for v* oscillograms (exactly as shown on v-oscillograms)
                try:
                    if (it.get("prefix") or "").lower() == "v":
                        tb_lines = build_textbox_lines(meta, file_prefix="v")
                        row = _parse_textbox_lines_to_row(tb_lines)
                        # Ensure No./Shape are always present
                        if "No." not in row:
                            row["No."] = str(it.get("No", ""))
                        if "Shape" not in row:
                            row["Shape"] = normalize_shape(str(meta.get("Shape", "")))
                        peaks_table_rows.append(row)
                except Exception:
                    pass

                images.append(out_png)
            try:
                insert_images_into_template_strict(
                    template, images, out_path,
                    width_cm=width_cm, height_cm=height_cm, lock_aspect=lock_aspect,
                    peaks_table_rows=peaks_table_rows,
                )
                messagebox.showinfo("Done", f"Generated successfully:\n{out_path}")
            except ValueError:
                parts, capacity = export_strict_template_multipart(
                    template, images, out_path,
                    width_cm=width_cm, height_cm=height_cm, lock_aspect=lock_aspect
                )
                # Append textbox table to the LAST document part (so it is "at the end")
                try:
                    if peaks_table_rows and parts:
                        dlast = Document(parts[-1])
                        _append_textbox_values_table(dlast, peaks_table_rows)
                        dlast.save(parts[-1])
                except Exception:
                    pass
                messagebox.showinfo(
                    "Done",
                    f"Template capacity is {capacity} images.\nGenerated multiple documents:\n" + "\n".join(parts)
                )
        except Exception as e:
            messagebox.showerror("Error", str(e))
def main():
    root = tk.Tk()
    AppV3_3(root)
    root.mainloop()
if __name__ == "__main__":
    main()
